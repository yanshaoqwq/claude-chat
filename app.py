import hashlib
import io
import json
import os
import random
import socket
import subprocess
import sys
import threading
import time
import uuid

import requests
from flask import Flask, Response, jsonify, render_template, request, stream_with_context

DEFAULT_UPSTREAM_URL = "https://cheaprouter.org/v1/messages?beta=true"
DEVICE_ID_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".device_id")

BETA_FLAGS = "context-1m-2025-08-07,interleaved-thinking-2025-05-14,redact-thinking-2026-02-12,context-management-2025-06-27,prompt-caching-scope-2026-01-05,effort-2025-11-24"

SEARCH_TOOL = {
    "name": "web_search",
    "description": "搜索互联网获取最新信息。涉及实时数据、新闻、价格、天气、版本号或知识截止后的事件时调用。",
    "input_schema": {
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "简洁的搜索关键词"}
        },
        "required": ["query"],
    },
}

SEARCH_USAGE_INSTRUCTION = (
    "当你需要联网信息时，**直接调用 web_search 工具**，不要先输出 \"我去搜索一下\"、"
    "\"I'll search for ...\"、\"让我搜一下\" 这类预告或解释文字。先调用工具，再用工具结果作答。\n"
    "拿到搜索结果后，请用自然语言综合提炼信息直接回答用户问题，不要把回答写成纯链接列表或 URL 罗列。"
    "如果需要标注来源，可在回答末尾用 markdown 链接形式简短给出 1-3 条主要参考。"
)

MAX_RETRIES = 3
RETRY_DELAYS = [0.5, 1.5, 3.0]
MAX_TOOL_ITERATIONS = 6

app = Flask(__name__)


# ─── Device ID ───────────────────────────────────────────────────────────────

def _get_first_mac():
    try:
        net_dir = "/sys/class/net"
        for iface in sorted(os.listdir(net_dir)):
            if iface == "lo":
                continue
            try:
                with open(os.path.join(net_dir, iface, "address")) as f:
                    mac = f.read().strip()
                if mac and mac != "00:00:00:00:00:00":
                    return mac
            except OSError:
                continue
    except OSError:
        pass
    node = uuid.getnode()
    return ":".join(f"{(node >> i) & 0xff:02x}" for i in range(40, -8, -8))


def _resolve_device_id():
    if os.path.isfile(DEVICE_ID_FILE):
        try:
            val = open(DEVICE_ID_FILE, "r").read().strip()
            if len(val) == 64:
                int(val, 16)
                return val
        except (IOError, ValueError):
            pass

    raw = f"{socket.gethostname()}:{_get_first_mac()}".encode()
    device_id = hashlib.sha256(raw).hexdigest()

    try:
        with open(DEVICE_ID_FILE, "w") as f:
            f.write(device_id)
        os.chmod(DEVICE_ID_FILE, 0o600)
    except OSError as e:
        print(f"[warn] cannot persist device_id: {e}", file=sys.stderr)

    return device_id


DEVICE_ID = _resolve_device_id()


# ─── Upstream request building ───────────────────────────────────────────────

def build_headers(api_key, session_id):
    return {
        "Accept": "application/json",
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "User-Agent": "claude-cli/2.1.145 (external, cli)",
        "X-Claude-Code-Session-Id": session_id,
        "X-Stainless-Arch": "arm64",
        "X-Stainless-Lang": "js",
        "X-Stainless-OS": "Linux",
        "X-Stainless-Package-Version": "0.94.0",
        "X-Stainless-Retry-Count": "0",
        "X-Stainless-Runtime": "node",
        "X-Stainless-Runtime-Version": "v24.3.0",
        "X-Stainless-Timeout": "600",
        "anthropic-beta": BETA_FLAGS,
        "anthropic-dangerous-direct-browser-access": "true",
        "anthropic-version": "2023-06-01",
        "x-app": "cli",
        "Connection": "keep-alive",
        "Accept-Encoding": "gzip, deflate, br, zstd",
    }


def build_payload(system_text, api_messages, model, max_tokens, effort, tools=None):
    payload = {
        "model": model,
        "system": [{"type": "text", "text": system_text or "", "cache_control": {"type": "ephemeral"}}],
        "messages": api_messages,
        "max_tokens": max_tokens,
        "thinking": {"type": "adaptive"},
        "context_management": {"edits": [{"type": "clear_thinking_20251015", "keep": "all"}]},
        "output_config": {"effort": effort},
        "stream": True,
    }
    if tools:
        payload["tools"] = tools
    return payload


def messages_to_api(messages):
    out = []
    for m in messages:
        role = m.get("role")
        if role not in ("user", "assistant"):
            continue
        content = m.get("content", "") or ""
        attachments = m.get("attachments") or []
        blocks = []
        for a in attachments:
            kind = a.get("kind")
            name = a.get("name") or ""
            if kind == "pdf":
                mode = a.get("mode") or "base64"
                if mode == "text":
                    txt = a.get("extracted_text") or ""
                    if not txt:
                        continue
                    blocks.append({
                        "type": "document",
                        "source": {"type": "text", "media_type": "text/plain", "data": txt},
                        "title": name,
                    })
                    continue
                data = a.get("data")
                if not data:
                    continue
                blocks.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": a.get("media_type") or "application/pdf",
                        "data": data,
                    },
                    "title": name,
                })
            elif kind == "image":
                data = a.get("data")
                if not data:
                    continue
                blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": a.get("media_type") or "image/png",
                        "data": data,
                    },
                })
            elif kind == "text":
                txt = a.get("text") or ""
                if not txt:
                    continue
                blocks.append({
                    "type": "document",
                    "source": {"type": "text", "media_type": "text/plain", "data": txt},
                    "title": name,
                })
        if content:
            blocks.append({"type": "text", "text": content})
        if not blocks:
            continue
        out.append({"role": role, "content": blocks})
    return out


def sse(obj):
    return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"


# ─── Web search API ──────────────────────────────────────────────────────────

def call_search_api(search_url, search_key, query, lang="zh"):
    if not search_url or not search_key:
        return None, "搜索 API 未配置"
    try:
        r = requests.post(
            search_url,
            headers={"Authorization": f"Bearer {search_key}", "Content-Type": "application/json"},
            json={"query": query, "lang": lang},
            timeout=60,
        )
        if r.status_code != 200:
            return None, f"搜索 HTTP {r.status_code}: {r.text[:300]}"
        data = r.json()
        if data.get("code") != 200:
            return None, f"搜索返回错误: {data.get('msg', 'unknown')}"
        return data.get("data"), None
    except requests.exceptions.RequestException as e:
        return None, f"搜索请求失败: {e}"
    except (ValueError, KeyError) as e:
        return None, f"搜索响应解析失败: {e}"


def format_search_result(data):
    if data is None:
        return "搜索未返回结果"
    if isinstance(data, str):
        return data
    if isinstance(data, dict):
        if "content" in data:
            return data["content"]
        if "answer" in data:
            return data["answer"]
        if "results" in data and isinstance(data["results"], list):
            parts = []
            for r in data["results"][:10]:
                title = r.get("title", "")
                url = r.get("url", "")
                snippet = r.get("snippet") or r.get("content") or ""
                parts.append(f"【{title}】{snippet}\n{url}")
            return "\n\n".join(parts)
        return json.dumps(data, ensure_ascii=False)[:4000]
    if isinstance(data, list):
        return "\n\n".join(str(x) for x in data)[:4000]
    return str(data)[:4000]


# ─── Upstream streaming with retry and tool_use loop ─────────────────────────

class StreamEvent:
    __slots__ = ("kind", "data")
    def __init__(self, kind, data):
        self.kind = kind
        self.data = data


def _stream_one_request(api_key, upstream_url, system_text, api_messages, model,
                       max_tokens, effort, tools, on_event):
    """Single upstream request. Calls on_event(StreamEvent) for each event.
    Returns (stop_reason, content_blocks, error_or_None, retryable_bool).
    content_blocks: full assistant message content for tool loop continuation.
    """
    session_id = str(uuid.uuid4())
    headers = build_headers(api_key, session_id)
    payload = build_payload(system_text, api_messages, model, max_tokens, effort, tools)
    payload["metadata"] = {
        "user_id": json.dumps({
            "device_id": DEVICE_ID,
            "account_uuid": "",
            "session_id": session_id,
        })
    }

    try:
        r = requests.post(upstream_url, headers=headers, json=payload, stream=True, timeout=600)
    except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
        return None, None, f"连接失败: {e}", True
    except requests.exceptions.RequestException as e:
        return None, None, f"请求失败: {e}", False

    with r:
        r.encoding = "utf-8"
        if r.status_code != 200:
            err = f"HTTP {r.status_code}: {r.text[:800]}"
            retryable = r.status_code >= 500 or r.status_code == 429
            return None, None, err, retryable

        content_blocks = []
        block_state = {}
        stop_reason = None

        try:
            for raw in r.iter_lines(decode_unicode=True, delimiter="\n"):
                if not raw or not raw.startswith("data: "):
                    continue
                data_str = raw[6:]
                if data_str.strip() == "[DONE]":
                    break
                try:
                    evt = json.loads(data_str)
                except json.JSONDecodeError:
                    continue

                etype = evt.get("type")

                if etype == "content_block_start":
                    idx = evt.get("index", 0)
                    block = evt.get("content_block", {})
                    btype = block.get("type")
                    block_state[idx] = {"type": btype, "data": dict(block)}
                    if btype == "tool_use":
                        block_state[idx]["input_json"] = ""
                        on_event(StreamEvent("tool_use_start", {
                            "tool_id": block.get("id"),
                            "name": block.get("name"),
                        }))

                elif etype == "content_block_delta":
                    idx = evt.get("index", 0)
                    d = evt.get("delta") or {}
                    dtype = d.get("type")
                    state = block_state.get(idx) or {}
                    if dtype == "text_delta":
                        chunk = d.get("text") or ""
                        state.setdefault("text", "")
                        state["text"] = state.get("text", "") + chunk
                        on_event(StreamEvent("text_delta", chunk))
                    elif dtype == "thinking_delta":
                        chunk = d.get("thinking") or ""
                        state.setdefault("thinking", "")
                        state["thinking"] = state.get("thinking", "") + chunk
                        on_event(StreamEvent("thinking_delta", chunk))
                    elif dtype == "input_json_delta":
                        chunk = d.get("partial_json") or ""
                        state["input_json"] = state.get("input_json", "") + chunk

                elif etype == "content_block_stop":
                    idx = evt.get("index", 0)
                    state = block_state.get(idx)
                    if not state:
                        continue
                    btype = state.get("type")
                    if btype == "text":
                        content_blocks.append({"type": "text", "text": state.get("text", "")})
                    elif btype == "thinking":
                        content_blocks.append({
                            "type": "thinking",
                            "thinking": state.get("thinking", ""),
                            "signature": state["data"].get("signature", ""),
                        })
                    elif btype == "tool_use":
                        try:
                            tool_input = json.loads(state.get("input_json") or "{}")
                        except json.JSONDecodeError:
                            tool_input = {}
                        content_blocks.append({
                            "type": "tool_use",
                            "id": state["data"].get("id"),
                            "name": state["data"].get("name"),
                            "input": tool_input,
                        })
                        on_event(StreamEvent("tool_use_input", {
                            "tool_id": state["data"].get("id"),
                            "input": tool_input,
                        }))

                elif etype == "message_delta":
                    delta = evt.get("delta") or {}
                    if "stop_reason" in delta:
                        stop_reason = delta["stop_reason"]

                elif etype == "error":
                    msg = (evt.get("error") or {}).get("message") or evt.get("message") or "upstream error"
                    return stop_reason, content_blocks, msg, True

        except (requests.exceptions.ChunkedEncodingError, requests.exceptions.ConnectionError) as e:
            return stop_reason, content_blocks, f"流中断: {e}", True

    return stop_reason, content_blocks, None, False


def stream_with_retry(api_key, upstream_url, system_text, api_messages, model,
                     max_tokens, effort, tools, on_event):
    """Wrap _stream_one_request with retry. Refuses to retry once any text/thinking has streamed."""
    streamed = {"v": False}

    def wrapped(evt):
        if evt.kind in ("text_delta", "thinking_delta"):
            streamed["v"] = True
        on_event(evt)

    last_err = None
    for attempt in range(MAX_RETRIES):
        stop_reason, blocks, err, retryable = _stream_one_request(
            api_key, upstream_url, system_text, api_messages, model,
            max_tokens, effort, tools, wrapped,
        )
        if err is None:
            return stop_reason, blocks, None
        last_err = err
        if not retryable or streamed["v"] or attempt == MAX_RETRIES - 1:
            return stop_reason, blocks, err
        delay = RETRY_DELAYS[attempt] * (0.8 + 0.4 * random.random())
        time.sleep(delay)
    return None, None, last_err


def run_chat_with_tools(api_key, upstream_url, system_text, api_messages, model,
                       max_tokens, effort, web_search_enabled, search_url, search_key,
                       on_event):
    """Run upstream loop, handling web_search tool calls. on_event yields events to client."""
    messages = list(api_messages)
    tools = [SEARCH_TOOL] if web_search_enabled else None

    for iteration in range(MAX_TOOL_ITERATIONS):
        stop_reason, blocks, err = stream_with_retry(
            api_key, upstream_url, system_text, messages, model,
            max_tokens, effort, tools, on_event,
        )
        if err:
            on_event(StreamEvent("error", err))
            return

        if stop_reason != "tool_use":
            return

        tool_uses = [b for b in (blocks or []) if b.get("type") == "tool_use"]
        if not tool_uses:
            return

        messages.append({"role": "assistant", "content": blocks})

        tool_results = []
        for tu in tool_uses:
            if tu.get("name") != "web_search":
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tu.get("id"),
                    "content": "未知工具",
                    "is_error": True,
                })
                continue
            query = (tu.get("input") or {}).get("query", "")
            data, search_err = call_search_api(search_url, search_key, query)
            if search_err:
                on_event(StreamEvent("tool_result", {
                    "tool_id": tu.get("id"),
                    "error": search_err,
                }))
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tu.get("id"),
                    "content": f"搜索失败：{search_err}。请基于已有知识作答并说明无法获取最新信息。",
                    "is_error": True,
                })
            else:
                formatted = format_search_result(data)
                on_event(StreamEvent("tool_result", {
                    "tool_id": tu.get("id"),
                    "result": {"content": formatted[:4000] + ("\n\n…(已截断)" if len(formatted) > 4000 else "")},
                }))
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tu.get("id"),
                    "content": formatted,
                })

        messages.append({"role": "user", "content": tool_results})

    on_event(StreamEvent("error", f"工具调用超过 {MAX_TOOL_ITERATIONS} 轮，已中止"))


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/extract", methods=["POST"])
def extract():
    import base64
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "no file"}), 400
    name = f.filename or "uploaded"
    raw = f.read()
    if len(raw) > 20 * 1024 * 1024:
        return jsonify({"ok": False, "error": "文件过大（最大 20MB）"}), 400
    lower = name.lower()
    image_exts = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp",
    }
    img_match = next((mt for ext, mt in image_exts.items() if lower.endswith(ext)), None)
    try:
        if lower.endswith(".pdf"):
            extracted_text, pages = _extract_pdf_text(raw)
            return jsonify({
                "ok": True, "name": name, "kind": "pdf",
                "media_type": "application/pdf",
                "data": base64.b64encode(raw).decode("ascii"),
                "size": len(raw),
                "extracted_text": extracted_text,
                "extracted_chars": len(extracted_text or ""),
                "pages": pages,
            })
        elif img_match:
            return jsonify({
                "ok": True, "name": name, "kind": "image",
                "media_type": img_match,
                "data": base64.b64encode(raw).decode("ascii"),
                "size": len(raw),
            })
        elif lower.endswith(".docx"):
            text = _extract_docx(raw)
        elif lower.endswith((".txt", ".md", ".markdown", ".log", ".csv", ".json", ".py", ".js", ".ts", ".html", ".css", ".yaml", ".yml", ".xml")):
            text = raw.decode("utf-8", errors="replace")
        else:
            return jsonify({"ok": False, "error": f"不支持的格式: {lower.rsplit('.', 1)[-1]}"}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": f"解析失败: {e}"}), 400

    text = (text or "").strip()
    if not text:
        return jsonify({"ok": False, "error": "未能从文件中提取出文本"}), 400
    if len(text) > 200_000:
        text = text[:200_000] + "\n\n…(已截断)"
    return jsonify({
        "ok": True, "name": name, "kind": "text",
        "text": text, "chars": len(text), "size": len(raw),
    })


def _extract_pdf_text(raw):
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", 0
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = len(reader.pages)
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            t = t.strip()
            if t:
                parts.append(f"=== Page {i + 1} ===\n{t}")
        text = "\n\n".join(parts)
        if len(text) > 200_000:
            text = text[:200_000] + "\n\n…(已截断)"
        return text, pages
    except Exception:
        return "", 0


def _extract_docx(raw):
    import docx as docx_mod
    doc = docx_mod.Document(io.BytesIO(raw))
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@app.route("/api/system/shutdown", methods=["POST"])
def sys_shutdown():
    threading.Thread(target=_delayed_exit, args=(0.4, 0), daemon=True).start()
    return jsonify({"ok": True, "action": "shutdown"})


@app.route("/api/system/restart", methods=["POST"])
def sys_restart():
    threading.Thread(target=_delayed_restart, args=(0.4,), daemon=True).start()
    return jsonify({"ok": True, "action": "restart"})


def _delayed_exit(delay, code):
    time.sleep(delay)
    os._exit(code)


def _delayed_restart(delay):
    time.sleep(delay)
    os._exit(42)


def _supervise():
    env = dict(os.environ)
    env["CHATUI_CHILD"] = "1"
    while True:
        try:
            proc = subprocess.Popen([sys.executable, *sys.argv], env=env)
        except Exception as e:
            print(f"[supervisor] failed to spawn: {e}", file=sys.stderr)
            sys.exit(1)
        try:
            code = proc.wait()
        except KeyboardInterrupt:
            proc.terminate()
            try:
                proc.wait(timeout=3)
            except subprocess.TimeoutExpired:
                proc.kill()
            sys.exit(0)
        if code == 42:
            print("[supervisor] child requested restart, respawning…", flush=True)
            time.sleep(0.5)
            continue
        sys.exit(code)


@app.route("/api/chat", methods=["POST"])
def chat():
    body = request.get_json(force=True) or {}
    api_key = (body.get("api_key") or "").strip()
    upstream_url = (body.get("upstream_url") or "").strip() or DEFAULT_UPSTREAM_URL
    system_text = body.get("system", "")
    messages = body.get("messages", [])
    model = body.get("model") or "claude-opus-4-7"
    max_tokens = int(body.get("max_tokens") or 64000)
    effort = body.get("effort") or "xhigh"
    web_search_enabled = bool(body.get("web_search_enabled"))
    search_url = (body.get("search_url") or "").strip()
    search_key = (body.get("search_key") or "").strip()

    if not api_key:
        return Response(
            sse({"type": "error", "message": "请先填写 API Key"}) + "data: [DONE]\n\n",
            mimetype="text/event-stream",
            status=400,
        )
    if web_search_enabled and (not search_url or not search_key):
        return Response(
            sse({"type": "error", "message": "已启用联网搜索，但未配置搜索 URL 或 Key"}) + "data: [DONE]\n\n",
            mimetype="text/event-stream",
            status=400,
        )

    if web_search_enabled:
        base = (system_text or "").strip()
        system_text = (base + "\n\n" + SEARCH_USAGE_INSTRUCTION) if base else SEARCH_USAGE_INSTRUCTION

    api_messages = messages_to_api(messages)

    queue = []
    queue_lock = threading.Lock()
    done = threading.Event()

    def push_event(evt):
        with queue_lock:
            queue.append(evt)

    def runner():
        try:
            run_chat_with_tools(
                api_key, upstream_url, system_text, api_messages, model,
                max_tokens, effort, web_search_enabled, search_url, search_key,
                push_event,
            )
        except Exception as e:
            push_event(StreamEvent("error", f"内部错误: {e}"))
        finally:
            done.set()

    def generate():
        threading.Thread(target=runner, daemon=True).start()
        try:
            while True:
                with queue_lock:
                    pending = queue[:]
                    queue.clear()
                for evt in pending:
                    yield _event_to_sse(evt)
                if done.is_set():
                    with queue_lock:
                        pending = queue[:]
                        queue.clear()
                    for evt in pending:
                        yield _event_to_sse(evt)
                    break
                time.sleep(0.03)
        finally:
            yield "data: [DONE]\n\n"

    resp = Response(stream_with_context(generate()), mimetype="text/event-stream")
    resp.headers["Cache-Control"] = "no-cache"
    resp.headers["X-Accel-Buffering"] = "no"
    return resp


def _event_to_sse(evt):
    if evt.kind == "text_delta":
        return sse({"type": "content_block_delta", "delta": {"type": "text_delta", "text": evt.data}})
    if evt.kind == "thinking_delta":
        return sse({"type": "content_block_delta", "delta": {"type": "thinking_delta", "thinking": evt.data}})
    if evt.kind == "tool_use_start":
        return sse({"type": "tool_use_start", **evt.data})
    if evt.kind == "tool_use_input":
        return sse({"type": "tool_use_input", **evt.data})
    if evt.kind == "tool_result":
        return sse({"type": "tool_result", **evt.data})
    if evt.kind == "error":
        return sse({"type": "error", "message": evt.data})
    return ""


if __name__ == "__main__":
    if os.environ.get("CHATUI_CHILD") == "1":
        app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
    else:
        _supervise()
